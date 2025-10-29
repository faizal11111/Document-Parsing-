from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pathlib import Path
from PIL import Image
import docx, pdfplumber, openpyxl, pytesseract, os, shutil
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploaded_files"
Path(UPLOAD_DIR).mkdir(exist_ok=True)

def parse_pdf(file_path):
    results = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables()
            results.append({
                "page_number": i + 1,
                "text_content": text,
                "tables": tables,
                "extraction_confidence": 1.0,
            })
    return results

def parse_image(file_path):
    img = Image.open(file_path)
    text = pytesseract.image_to_string(img)
    return [{
        "page_number": 1,
        "text_content": text,
        "tables": [],
        "extraction_confidence": "N/A"
    }]

def parse_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    tables = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            table_rows.append([cell.text for cell in row.cells])
        tables.append(table_rows)
    return [{
        "page_number": 1,
        "text_content": text,
        "tables": tables,
        "extraction_confidence": 1.0
    }]

def parse_excel(file_path):
    wb = openpyxl.load_workbook(file_path)
    data = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(cell) if cell is not None else "" for cell in row])
        data.append({"sheet_name": sheet, "rows": rows})
    return [{
        "page_number": 1,
        "text_content": "",
        "tables": data,
        "extraction_confidence": 1.0
    }]

def parse_txt(file_path):
    with open(file_path, "r", encoding='utf8') as f:
        text = f.read()
    return [{
        "page_number": 1,
        "text_content": text,
        "tables": [],
        "extraction_confidence": 1.0
    }]
@app.post("/upload/")
async def upload_files(files: list[UploadFile] = File(...)):
    responses = []
    for file in files:
        file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_location, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)   
        ext = Path(file_location).suffix.lower()
        parsed_output = None
        try:
            if ext == ".pdf":
                parsed_output = parse_pdf(file_location)
            elif ext in [".jpg", ".jpeg", ".png"]:
                parsed_output = parse_image(file_location)
            elif ext == ".docx":
                parsed_output = parse_docx(file_location)
            elif ext in [".xls", ".xlsx"]:
                parsed_output = parse_excel(file_location)
            elif ext == ".txt":
                parsed_output = parse_txt(file_location)
            else:
                parsed_output = {"error": "Unsupported file type"}
        except Exception as e:
            parsed_output = {"error": f"Parsing failed: {str(e)}"}
        responses.append({
            "filename": file.filename,
            "parsed_output": parsed_output  
        })
    return {"files": responses}
